import gradio as gr
from faster_whisper import WhisperModel
from docx import Document
import os
import time
import logging

# Setup logging
logging.basicConfig(filename="transcribe.log", level=logging.DEBUG, format="%(asctime)s - %(levelname)s - %(message)s")

# Load model dynamically based on user selection
def load_model(size, device="cuda", compute_type="float16"):
    logging.debug(f"Loading model: size={size}, device={device}, compute_type={compute_type}")
    return WhisperModel(size, device=device, compute_type=compute_type)

# Transcription logic
def transcribe(file_obj, model_size, language, include_timestamps):
    try:
        if not file_obj:
            logging.warning("No file uploaded.")
            return "No file uploaded.", None

        audio_path = file_obj.name
        logging.debug(f"Received file: {audio_path}")

        model = load_model(model_size)

        logging.debug(f"Starting transcription: model_size={model_size}, language={language}, timestamps={include_timestamps}")
        segments, info = model.transcribe(
            audio_path,
            language=None if language == "Auto" else language,
            vad_filter=True,
            vad_parameters={"threshold": 0.5}
        )

        logging.debug(f"Transcription metadata: {info}")

        # Format transcript
        lines = []
        for seg in segments:
            if include_timestamps:
                lines.append(f"[{seg.start:.2f}s - {seg.end:.2f}s] {seg.text}")
            else:
                lines.append(seg.text)

        transcript = "\n".join(lines)

        # Save to Word document
        timestamp = time.strftime("%Y%m%d-%H%M%S")
        doc_path = f"transcript_{timestamp}.docx"
        doc = Document()
        doc.add_heading("Transcription", level=1)
        doc.add_paragraph(transcript)
        doc.save(doc_path)

        logging.info(f"Transcription complete. Saved to {doc_path}")
        return transcript, doc_path

    except Exception as e:
        logging.error(f"Error during transcription: {str(e)}", exc_info=True)
        return f"Error during transcription: {str(e)}", None

# Gradio UI
with gr.Blocks(title="FasterWhisper Transcription Tool") as demo:
    gr.Markdown("## 🎙️ FasterWhisper Audio Transcription")
    gr.Markdown("Upload an audio file, choose your model and language, and get a Word document with the transcription.")

    with gr.Row():
        audio_input = gr.File(label="Upload Audio File (.wav, .mp3, etc.)")
        model_size = gr.Dropdown(choices=["tiny", "base", "medium", "large-v2"], value="base", label="Model Size")
        language = gr.Dropdown(choices=["Auto", "en", "zh", "es", "fr", "de", "ja", "ko"], value="Auto", label="Language")
        include_timestamps = gr.Checkbox(label="Include Timestamps", value=True)

    transcribe_btn = gr.Button("Transcribe")

    transcript_output = gr.Textbox(label="Transcription", lines=10)
    docx_output = gr.File(label="Download Word Document")

    transcribe_btn.click(
        transcribe,
        inputs=[audio_input, model_size, language, include_timestamps],
        outputs=[transcript_output, docx_output]
    )

if __name__ == "__main__":
    demo.launch(debug=True)
